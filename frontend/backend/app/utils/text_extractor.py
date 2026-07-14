"""
Resume text extraction — PDF (pdfplumber) and DOCX (python-docx).

Why this file exists
---------------------
Text extraction is pure, stateless, format-specific logic: given a file
path and a known type, produce plain text. Isolating it here (rather than
inline in the service) makes it independently unit-testable and reusable
if another module later needs to extract text from a PDF/DOCX (e.g. a
LinkedIn PDF export in the LinkedIn Optimizer module reuses
`extract_text_from_pdf` directly).

How it works
------------
- `extract_text_from_pdf` opens the file with pdfplumber and concatenates
  each page's extracted text.
- `extract_text_from_docx` opens the file with python-docx and joins all
  paragraph text (including table cell text, since resumes sometimes use
  tables for skills/experience layout).
- `extract_text` is the single entrypoint the service calls; it dispatches
  by extension and wraps any parsing failure in `ResumeParsingError` so
  the caller doesn't need to know about pdfplumber/python-docx internals
  or their specific exception types.

Both extraction functions are synchronous/blocking (pdfplumber and
python-docx do blocking file I/O and CPU-bound parsing) — the service
layer runs them via `asyncio.to_thread` so the event loop isn't blocked.

Where future code should go
----------------------------
If a new resume format is supported later (e.g. .txt or .rtf), add a
sibling `extract_text_from_<format>` function and a new branch in
`extract_text`.
"""

import logging
from pathlib import Path

import docx
import pdfplumber

from app.core.exceptions import ResumeParsingError

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract all text from a PDF file, page by page, using pdfplumber."""
    try:
        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as exc:  # pdfplumber/pdfminer can raise various internal errors
        logger.warning("Failed to extract text from PDF '%s': %s", file_path.name, exc)
        raise ResumeParsingError(
            "Could not read the PDF file. It may be corrupted, password-protected, "
            "or a scanned image without a text layer."
        ) from exc


def extract_text_from_docx(file_path: Path) -> str:
    """Extract all paragraph and table text from a DOCX file using python-docx."""
    try:
        document = docx.Document(str(file_path))

        text_parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

        # Resumes frequently lay out skills/experience in tables — include
        # table cell text so it isn't silently dropped.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        return "\n".join(text_parts).strip()
    except Exception as exc:  # python-docx raises various zipfile/XML errors on corrupt files
        logger.warning("Failed to extract text from DOCX '%s': %s", file_path.name, exc)
        raise ResumeParsingError(
            "Could not read the DOCX file. It may be corrupted or not a valid Word document."
        ) from exc


def extract_text(file_path: Path, extension: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.

    Raises ResumeParsingError if the extension is unrecognized (should be
    unreachable if upstream validation ran first) or if extraction yields
    no usable text.
    """
    if extension == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif extension == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        raise ResumeParsingError(f"Unsupported extension '{extension}' for text extraction.")

    if not text:
        raise ResumeParsingError(
            "No extractable text was found in the resume. It may be a scanned "
            "image, empty, or use an unsupported layout."
        )

    return text
